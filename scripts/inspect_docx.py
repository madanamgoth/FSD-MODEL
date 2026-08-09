"""Quick inspection of sample FSD DOCX files."""
from pathlib import Path
from collections import Counter
from docx import Document

root = Path(r"c:\FSD MODEL\data\fsds\FSD")
word_dir = root / "FSD-WORD-DOCUMENT"
base_dir = root / "FSD-BASE"

files = sorted(word_dir.glob("*.docx"), key=lambda p: p.stat().st_size)
picks = []
if base_dir.exists():
    picks.extend(sorted(base_dir.glob("*.docx"))[:2])
if files:
    picks.append(files[len(files) // 4])
    picks.append(files[len(files) // 2])
    picks.append(files[-1])
for f in files:
    n = f.name.lower()
    if any(k in n for k in ("sms", "notification", "integration", "payment", "otp", "callback")):
        picks.append(f)
        if len([p for p in picks if "WORD" in str(p)]) >= 5:
            break

seen = set()
uniq = []
for p in picks:
    if p.exists() and p.resolve() not in seen:
        seen.add(p.resolve())
        uniq.append(p)

print(f"Inspecting {len(uniq)} files\n")
for path in uniq[:7]:
    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = doc.tables
    img_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_count += 1
    except Exception as exc:
        img_count = f"err:{exc}"

    styles = Counter()
    for p in doc.paragraphs:
        if p.text.strip() and p.style and p.style.name:
            styles[p.style.name] += 1

    print("=" * 80)
    print("FILE:", path.name)
    print("SIZE_KB:", round(path.stat().st_size / 1024, 1))
    print("paragraphs_nonempty:", len(paras))
    print("tables:", len(tables))
    print("images:", img_count)
    print("top_styles:", styles.most_common(8))
    print("--- first 8 paragraphs ---")
    for t in paras[:8]:
        print(" ", t[:180].replace("\n", " "))
    if tables:
        t0 = tables[0]
        rows = min(4, len(t0.rows))
        cols = min(5, len(t0.columns))
        print(f"--- table[0] sample {rows}x{cols} ---")
        for r in range(rows):
            cells = []
            for c in range(min(cols, len(t0.rows[r].cells))):
                cells.append(t0.rows[r].cells[c].text.strip().replace("\n", " ")[:35])
            print(" | ".join(cells))
    print()
